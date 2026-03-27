from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

# Create a new Document
doc = Document()

# Set up styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Read the markdown file
with open('NUTRITRACK_SDLC_DOCUMENTATION.md', 'r', encoding='utf-8') as f:
    content = f.read()

# Split content into lines
lines = content.split('\n')

i = 0
while i < len(lines):
    line = lines[i]
    
    # Skip empty lines
    if not line.strip():
        i += 1
        continue
    
    # Handle headers
    if line.startswith('# '):
        p = doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        p = doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        p = doc.add_heading(line[4:], level=3)
    elif line.startswith('#### '):
        p = doc.add_heading(line[5:], level=4)
    
    # Handle code blocks
    elif line.startswith('```'):
        code_lines = []
        i += 1
        while i < len(lines) and not lines[i].startswith('```'):
            code_lines.append(lines[i])
            i += 1
        
        # Add code block
        p = doc.add_paragraph()
        p.style = 'Normal'
        run = p.add_run('\n'.join(code_lines))
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Inches(0.5)
        
    # Handle tables
    elif line.startswith('|') and '|' in line:
        # Collect table rows
        table_rows = []
        while i < len(lines) and lines[i].startswith('|'):
            row = [cell.strip() for cell in lines[i].split('|')[1:-1]]
            table_rows.append(row)
            i += 1
            # Skip separator row
            if i < len(lines) and lines[i].startswith('|') and '---' in lines[i]:
                i += 1
        
        if table_rows:
            # Create table
            table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
            table.style = 'Light Grid Accent 1'
            
            for row_idx, row_data in enumerate(table_rows):
                for col_idx, cell_data in enumerate(row_data):
                    table.rows[row_idx].cells[col_idx].text = cell_data
        
        i -= 1  # Adjust because we'll increment at the end
    
    # Handle bullet points
    elif line.startswith('- ') or line.startswith('* '):
        doc.add_paragraph(line[2:], style='List Bullet')
    
    # Handle numbered lists
    elif re.match(r'^\d+\. ', line):
        text = re.sub(r'^\d+\. ', '', line)
        doc.add_paragraph(text, style='List Number')
    
    # Handle horizontal rules
    elif line.strip() == '---':
        doc.add_paragraph('_' * 80)
    
    # Handle bold text in paragraphs
    elif '**' in line:
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
    
    # Regular paragraph
    else:
        doc.add_paragraph(line)
    
    i += 1

# Save the document
doc.save('NUTRITRACK_SDLC_DOCUMENTATION.docx')
print("✅ Successfully converted to DOCX!")
print("📄 File saved as: NUTRITRACK_SDLC_DOCUMENTATION.docx")
